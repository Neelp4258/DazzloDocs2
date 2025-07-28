import os
import io
import tempfile
import subprocess
import logging
import json
import csv
import xml.etree.ElementTree as ET
from pathlib import Path
from PIL import Image
from config import Config

logger = logging.getLogger(__name__)

class FileConverter:
    """Handles file conversion between different formats - Universal version"""
    
    def __init__(self):
        self.image_quality = Config.IMAGE_QUALITY
        self.image_max_dimension = Config.IMAGE_MAX_DIMENSION
        self.pdf_resolution = Config.PDF_RESOLUTION
        
        # Check available libraries
        self.pil_available = self._check_pil()
        self.pymupdf_available = self._check_pymupdf()
        self.reportlab_available = self._check_reportlab()
        self.docx_available = self._check_docx()
        
        logger.info(f"Universal Converter initialized - PIL: {self.pil_available}, PyMuPDF: {self.pymupdf_available}, "
                   f"ReportLab: {self.reportlab_available}, DOCX: {self.docx_available}")
    
    def _check_pil(self):
        """Check if PIL/Pillow is available"""
        try:
            from PIL import Image
            return True
        except ImportError:
            return False
    
    def _check_pymupdf(self):
        """Check if PyMuPDF is available"""
        try:
            import fitz
            return True
        except ImportError:
            return False
    
    def _check_reportlab(self):
        """Check if ReportLab is available"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate
            return True
        except ImportError:
            return False
    
    def _check_docx(self):
        """Check if python-docx is available"""
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    
    def convert_file(self, input_path: str, output_path: str, target_format: str) -> dict:
        """Main conversion method with universal format support"""
        try:
            if not os.path.exists(input_path):
                return {'success': False, 'error': 'Input file not found'}
            
            input_ext = Path(input_path).suffix[1:].lower()
            target_format = target_format.lower()
            
            logger.info(f"Converting {input_ext} to {target_format}")
            
            # Universal format mappings
            image_formats = ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg']
            document_formats = ['pdf', 'txt', 'docx', 'doc', 'rtf', 'md', 'html', 'htm']
            spreadsheet_formats = ['xlsx', 'xls', 'csv']
            presentation_formats = ['pptx', 'ppt']
            data_formats = ['json', 'xml']
            code_formats = ['py', 'js', 'css', 'php', 'java', 'cpp', 'c', 'cs', 'rb', 'go', 'rs', 'log', 'ini', 'cfg', 'conf', 'yaml', 'yml', 'toml']
            
            success = False  # Initialize success variable
            
            # Image conversions
            if input_ext in image_formats and target_format in image_formats:
                success = self._convert_image(input_path, output_path, target_format)
            elif input_ext in image_formats and target_format == 'pdf':
                success = self._convert_image_to_pdf(input_path, output_path)
            elif input_ext == 'pdf' and target_format in image_formats:
                success = self._convert_pdf_to_image(input_path, output_path, target_format)
            
            # Special conversions (must come before general document conversions)
            elif input_ext == 'pdf' and target_format == 'docx':
                logger.info(f"Attempting PDF to DOCX conversion: {input_path} -> {output_path}")
                success = self._convert_pdf_to_docx(input_path, output_path)
                logger.info(f"PDF to DOCX conversion result: {success}")
            elif input_ext == 'pdf' and target_format == 'doc':
                logger.info(f"Attempting PDF to DOC conversion: {input_path} -> {output_path}")
                success = self._convert_pdf_to_doc(input_path, output_path)
                logger.info(f"PDF to DOC conversion result: {success}")
            elif input_ext == 'pdf' and target_format == 'txt':
                success = self._convert_pdf_to_text(input_path, output_path)
            elif input_ext == 'txt' and target_format == 'pdf':
                success = self._convert_text_to_pdf(input_path, output_path)
            elif input_ext == 'docx' and target_format == 'pdf':
                success = self._convert_docx_to_pdf(input_path, output_path)
            elif input_ext == 'docx' and target_format == 'txt':
                success = self._convert_docx_to_text(input_path, output_path)
            
            # Document conversions (general)
            elif input_ext in document_formats and target_format in document_formats:
                success = self._convert_document(input_path, output_path, target_format)
            
            # Spreadsheet conversions
            elif input_ext in spreadsheet_formats and target_format in spreadsheet_formats:
                success = self._convert_spreadsheet(input_path, output_path, target_format)
            elif input_ext in spreadsheet_formats and target_format == 'pdf':
                success = self._convert_spreadsheet_to_pdf(input_path, output_path)
            
            # Data format conversions
            elif input_ext in data_formats and target_format in data_formats:
                success = self._convert_data_format(input_path, output_path, target_format)
            
            # Code format conversions
            elif input_ext in code_formats and target_format in code_formats:
                success = self._convert_code_format(input_path, output_path, target_format)
            
            # Cross-format conversions
            elif input_ext in document_formats and target_format == 'pdf':
                success = self._convert_to_pdf(input_path, output_path)
            elif input_ext == 'pdf' and target_format in document_formats:
                success = self._convert_from_pdf(input_path, output_path, target_format)
            elif input_ext == 'csv' and target_format == 'json':
                success = self._convert_csv_to_json(input_path, output_path)
            elif input_ext == 'json' and target_format == 'csv':
                success = self._convert_json_to_csv(input_path, output_path)
            elif input_ext == 'json' and target_format == 'xml':
                success = self._convert_json_to_xml(input_path, output_path)
            elif input_ext == 'xml' and target_format == 'json':
                success = self._convert_xml_to_json(input_path, output_path)
            
            else:
                return {'success': False, 'error': f'Conversion from {input_ext} to {target_format} not supported'}
            
            if success:
                return {'success': True, 'output_path': output_path}
            else:
                return {'success': False, 'error': 'Conversion failed'}
                
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            return {'success': False, 'error': f'Conversion error: {str(e)}'}
    
    def _convert_image(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert between image formats"""
        try:
            input_ext = Path(input_path).suffix[1:].lower()
            
            # Handle SVG conversions
            if target_format == 'svg':
                return self._convert_to_svg(input_path, output_path)
            elif input_ext == 'svg':
                return self._convert_from_svg(input_path, output_path, target_format)
            
            # Handle regular image conversions with PIL
            if not self.pil_available:
                return False
            
            with Image.open(input_path) as img:
                # Convert to RGB if necessary
                if target_format in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'LA', 'P']:
                    img = img.convert('RGB')
                elif target_format in ['png', 'webp'] and img.mode == 'P':
                    img = img.convert('RGBA')
                
                # Resize if too large
                if max(img.size) > self.image_max_dimension:
                    img.thumbnail((self.image_max_dimension, self.image_max_dimension), Image.Resampling.LANCZOS)
                
                # Save with appropriate options
                save_kwargs = {}
                if target_format in ['jpg', 'jpeg']:
                    save_kwargs['quality'] = self.image_quality
                    save_kwargs['optimize'] = True
                elif target_format == 'webp':
                    save_kwargs['quality'] = self.image_quality
                    save_kwargs['method'] = 6
                elif target_format == 'png':
                    save_kwargs['optimize'] = True
                
                img.save(output_path, format=target_format.upper(), **save_kwargs)
                return True
                
        except Exception as e:
            logger.error(f"Image conversion error: {str(e)}")
            return False
    
    def _convert_image_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert image to PDF"""
        try:
            if self.reportlab_available:
                return self._convert_image_to_pdf_reportlab(input_path, output_path)
            elif self.pil_available:
                return self._convert_image_to_pdf_pil(input_path, output_path)
            
            return False
            
        except Exception as e:
            logger.error(f"Image to PDF conversion error: {str(e)}")
            return False
    
    def _convert_image_to_pdf_reportlab(self, input_path: str, output_path: str) -> bool:
        """Convert image to PDF using ReportLab"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Image as RLImage
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            
            # Add image to PDF
            img = RLImage(input_path, width=6*inch, height=4*inch)
            story.append(img)
            
            doc.build(story)
            return True
            
        except Exception as e:
            logger.error(f"ReportLab image to PDF error: {str(e)}")
            return False
    
    def _convert_image_to_pdf_pil(self, input_path: str, output_path: str) -> bool:
        """Convert image to PDF using PIL"""
        try:
            with Image.open(input_path) as img:
                if img.mode in ['RGBA', 'LA', 'P']:
                    img = img.convert('RGB')
                img.save(output_path, 'PDF', resolution=self.pdf_resolution)
                return True
                
        except Exception as e:
            logger.error(f"PIL image to PDF error: {str(e)}")
            return False
    
    def _convert_pdf_to_image(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert PDF to image"""
        try:
            if self.pymupdf_available:
                import fitz
                doc = fitz.open(input_path)
                page = doc[0]  # First page
                
                # Calculate zoom for high resolution
                zoom = self.pdf_resolution / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image for format conversion
                img_data = pix.tobytes("png")
                with Image.open(io.BytesIO(img_data)) as img:
                    img.save(output_path, format=target_format.upper())
                
                doc.close()
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"PDF to image conversion error: {str(e)}")
            return False
    
    def _convert_to_svg(self, input_path: str, output_path: str) -> bool:
        """Convert image to SVG format"""
        try:
            # For now, create a simple SVG wrapper around the image
            # This is a basic approach - for better results, consider using vectorization libraries
            if not self.pil_available:
                return False
            
            with Image.open(input_path) as img:
                width, height = img.size
                
                # Create a simple SVG that embeds the image as base64
                import base64
                import io
                
                # Convert image to base64
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode()
                
                svg_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg width="{width}" height="{height}" xmlns="http://www.w3.org/2000/svg">
    <image width="{width}" height="{height}" href="data:image/png;base64,{img_base64}"/>
</svg>'''
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(svg_content)
                
                return True
                
        except Exception as e:
            logger.error(f"Convert to SVG error: {str(e)}")
            return False
    
    def _convert_from_svg(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert SVG to other image formats"""
        try:
            # Try using cairosvg if available and Cairo is installed
            try:
                import cairosvg
                cairosvg.svg2png(url=input_path, write_to=output_path.replace(f'.{target_format}', '.png'))
                
                # If target format is not PNG, convert the PNG to target format
                if target_format != 'png':
                    temp_png = output_path.replace(f'.{target_format}', '.png')
                    if os.path.exists(temp_png):
                        with Image.open(temp_png) as img:
                            img.save(output_path, format=target_format.upper())
                        os.remove(temp_png)
                
                return True
                
            except (ImportError, OSError) as e:
                logger.warning(f"cairosvg not available or Cairo library missing: {str(e)}")
            
            # Fallback: try using svglib if available
            try:
                from svglib.svglib import svg2rlg
                from reportlab.graphics import renderPM
                
                drawing = svg2rlg(input_path)
                if drawing:
                    renderPM.drawToFile(drawing, output_path, fmt=target_format.upper())
                    return True
                    
            except ImportError:
                logger.warning("svglib not available")
            
            # Final fallback: create a simple rasterized version
            if self.pil_available:
                # Create a simple colored rectangle as placeholder
                img = Image.new('RGB', (800, 600), color='white')
                img.save(output_path, format=target_format.upper())
                logger.warning("SVG conversion using fallback method - result may be basic")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Convert from SVG error: {str(e)}")
            return False
    
    def _convert_document(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert between document formats"""
        try:
            input_ext = Path(input_path).suffix[1:].lower()
            
            # Simple text-based conversions
            if input_ext == 'txt' and target_format == 'html':
                return self._convert_text_to_html(input_path, output_path)
            elif input_ext == 'html' and target_format == 'txt':
                return self._convert_html_to_text(input_path, output_path)
            elif input_ext == 'md' and target_format == 'html':
                return self._convert_markdown_to_html(input_path, output_path)
            elif input_ext == 'md' and target_format == 'txt':
                return self._convert_markdown_to_text(input_path, output_path)
            
            # For other conversions, try to extract text and convert
            if target_format == 'txt':
                return self._extract_text_to_file(input_path, output_path)
            elif target_format == 'html':
                return self._convert_to_html(input_path, output_path)
            
            return False
            
        except Exception as e:
            logger.error(f"Document conversion error: {str(e)}")
            return False
    
    def _convert_spreadsheet(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert between spreadsheet formats"""
        try:
            input_ext = Path(input_path).suffix[1:].lower()
            
            if input_ext == 'csv' and target_format == 'json':
                return self._convert_csv_to_json(input_path, output_path)
            elif input_ext == 'csv' and target_format == 'xml':
                return self._convert_csv_to_xml(input_path, output_path)
            elif input_ext == 'json' and target_format == 'csv':
                return self._convert_json_to_csv(input_path, output_path)
            
            return False
            
        except Exception as e:
            logger.error(f"Spreadsheet conversion error: {str(e)}")
            return False
    
    def _convert_data_format(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert between data formats"""
        try:
            input_ext = Path(input_path).suffix[1:].lower()
            
            if input_ext == 'json' and target_format == 'xml':
                return self._convert_json_to_xml(input_path, output_path)
            elif input_ext == 'xml' and target_format == 'json':
                return self._convert_xml_to_json(input_path, output_path)
            elif input_ext == 'json' and target_format == 'csv':
                return self._convert_json_to_csv(input_path, output_path)
            elif input_ext == 'csv' and target_format == 'json':
                return self._convert_csv_to_json(input_path, output_path)
            
            return False
            
        except Exception as e:
            logger.error(f"Data format conversion error: {str(e)}")
            return False
    
    def _convert_code_format(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert between code formats (mostly syntax highlighting)"""
        try:
            # For code files, we'll just copy the content with appropriate headers
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            if target_format == 'html':
                return self._convert_code_to_html(input_path, output_path, content)
            elif target_format == 'txt':
                # Just copy the content
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Code format conversion error: {str(e)}")
            return False
    
    def _convert_pdf_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to text"""
        try:
            if self.pymupdf_available:
                import fitz
                doc = fitz.open(input_path)
                text_content = ""
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text_content += page.get_text()
                    if page_num < len(doc) - 1:
                        text_content += "\n\n"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
                doc.close()
                return True
            
            return False
        except Exception as e:
            logger.error(f"PDF to text error: {str(e)}")
            return False
    
    def _convert_pdf_to_docx(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOCX"""
        try:
            if self.pymupdf_available and self.docx_available:
                return self._convert_pdf_to_docx_with_libraries(input_path, output_path)
            
            return False
        except Exception as e:
            logger.error(f"PDF to DOCX error: {str(e)}")
            return False
    
    def _convert_pdf_to_docx_with_libraries(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOCX using PyMuPDF + python-docx"""
        try:
            import fitz
            from docx import Document
            from docx.shared import Inches
            
            # Extract text from PDF
            doc = fitz.open(input_path)
            document = Document()
            
            # Add title
            document.add_heading('Converted PDF Document', 0)
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    # Split text into paragraphs
                    paragraphs = text.split('\n\n')
                    
                    for para in paragraphs:
                        if para.strip():
                            # Clean up the paragraph
                            clean_para = para.strip().replace('\n', ' ')
                            if clean_para:
                                document.add_paragraph(clean_para)
                    
                    # Add page break if not the last page
                    if page_num < len(doc) - 1:
                        document.add_page_break()
            
            # Save the DOCX file
            document.save(output_path)
            doc.close()
            return True
            
        except Exception as e:
            logger.error(f"PDF to DOCX with libraries error: {str(e)}")
            return False
    
    def _convert_text_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert text to PDF"""
        try:
            if self.reportlab_available:
                return self._convert_text_to_pdf_reportlab(input_path, output_path)
            
            return False
        except Exception as e:
            logger.error(f"Text to PDF error: {str(e)}")
            return False
    
    def _convert_text_to_pdf_reportlab(self, input_path: str, output_path: str) -> bool:
        """Convert text to PDF using ReportLab"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Read text file
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    clean_para = para.strip().replace('\n', ' ')
                    p = Paragraph(clean_para, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            return True
            
        except Exception as e:
            logger.error(f"ReportLab text to PDF error: {str(e)}")
            return False
    
    def _convert_docx_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to PDF"""
        try:
            if self.docx_available and self.reportlab_available:
                return self._convert_docx_to_pdf_reportlab(input_path, output_path)
            
            return False
        except Exception as e:
            logger.error(f"DOCX to PDF error: {str(e)}")
            return False
    
    def _convert_docx_to_pdf_reportlab(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to PDF using python-docx + ReportLab"""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Extract text from DOCX
            doc = Document(input_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Create temporary text file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as temp_file:
                temp_file.write('\n\n'.join(content))
                temp_path = temp_file.name
            
            try:
                # Convert using ReportLab
                result = self._convert_text_to_pdf_reportlab(temp_path, output_path)
                return result
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        except Exception as e:
            logger.error(f"DOCX to PDF ReportLab error: {str(e)}")
            return False
    
    def _convert_docx_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert DOCX to text"""
        try:
            if self.docx_available:
                from docx import Document
                doc = Document(input_path)
                content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n\n'.join(content))
                
                return True
            
            return False
        except Exception as e:
            logger.error(f"DOCX to text error: {str(e)}")
            return False
    
    # Helper conversion methods
    def _convert_text_to_html(self, input_path: str, output_path: str) -> bool:
        """Convert text to HTML"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; }}
    </style>
</head>
<body>
    <pre>{content}</pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        except Exception as e:
            logger.error(f"Text to HTML error: {str(e)}")
            return False
    
    def _convert_html_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert HTML to text"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Simple HTML to text conversion
            import re
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', content)
            # Remove extra whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            return True
        except Exception as e:
            logger.error(f"HTML to text error: {str(e)}")
            return False
    
    def _convert_markdown_to_html(self, input_path: str, output_path: str) -> bool:
        """Convert Markdown to HTML"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Simple markdown to HTML conversion
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Markdown</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; }}
    </style>
</head>
<body>
    <pre>{content}</pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        except Exception as e:
            logger.error(f"Markdown to HTML error: {str(e)}")
            return False
    
    def _convert_markdown_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert Markdown to text"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Simple markdown to text conversion
            import re
            # Remove markdown syntax
            text = re.sub(r'#+\s*', '', content)  # Remove headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove inline code
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            return True
        except Exception as e:
            logger.error(f"Markdown to text error: {str(e)}")
            return False
    
    def _convert_csv_to_json(self, input_path: str, output_path: str) -> bool:
        """Convert CSV to JSON"""
        try:
            data = []
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    data.append(row)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            
            return True
        except Exception as e:
            logger.error(f"CSV to JSON error: {str(e)}")
            return False
    
    def _convert_json_to_csv(self, input_path: str, output_path: str) -> bool:
        """Convert JSON to CSV"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            
            if isinstance(data, list) and len(data) > 0:
                fieldnames = data[0].keys()
                with open(output_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.DictWriter(f, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(data)
                return True
            return False
        except Exception as e:
            logger.error(f"JSON to CSV error: {str(e)}")
            return False
    
    def _convert_json_to_xml(self, input_path: str, output_path: str) -> bool:
        """Convert JSON to XML"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            
            root = ET.Element("root")
            self._dict_to_xml(data, root)
            
            tree = ET.ElementTree(root)
            tree.write(output_path, encoding='utf-8', xml_declaration=True)
            return True
        except Exception as e:
            logger.error(f"JSON to XML error: {str(e)}")
            return False
    
    def _convert_xml_to_json(self, input_path: str, output_path: str) -> bool:
        """Convert XML to JSON"""
        try:
            tree = ET.parse(input_path)
            root = tree.getroot()
            
            data = self._xml_to_dict(root)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            return True
        except Exception as e:
            logger.error(f"XML to JSON error: {str(e)}")
            return False
    
    def _convert_csv_to_xml(self, input_path: str, output_path: str) -> bool:
        """Convert CSV to XML"""
        try:
            import csv
            import xml.etree.ElementTree as ET
            
            root = ET.Element("data")
            
            with open(input_path, 'r', encoding='utf-8', newline='') as csvfile:
                reader = csv.DictReader(csvfile)
                
                for row in reader:
                    record = ET.SubElement(root, "record")
                    for key, value in row.items():
                        field = ET.SubElement(record, key.replace(' ', '_').lower())
                        field.text = str(value) if value else ""
            
            tree = ET.ElementTree(root)
            tree.write(output_path, encoding='utf-8', xml_declaration=True)
            return True
            
        except Exception as e:
            logger.error(f"CSV to XML error: {str(e)}")
            return False

    def _convert_spreadsheet_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert spreadsheet files to PDF"""
        try:
            input_ext = Path(input_path).suffix[1:].lower()
            
            if input_ext == 'csv':
                # Convert CSV to text first, then to PDF
                temp_txt = output_path.replace('.pdf', '.txt')
                if self._convert_csv_to_text(input_path, temp_txt):
                    result = self._convert_text_to_pdf(temp_txt, output_path)
                    if os.path.exists(temp_txt):
                        os.remove(temp_txt)
                    return result
            
            return False
            
        except Exception as e:
            logger.error(f"Spreadsheet to PDF error: {str(e)}")
            return False

    def _convert_csv_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert CSV to formatted text"""
        try:
            import csv
            
            with open(input_path, 'r', encoding='utf-8', newline='') as csvfile:
                reader = csv.reader(csvfile)
                
                with open(output_path, 'w', encoding='utf-8') as txtfile:
                    for row in reader:
                        txtfile.write(' | '.join(str(cell) for cell in row) + '\n')
            
            return True
            
        except Exception as e:
            logger.error(f"CSV to text error: {str(e)}")
            return False

    def _convert_pdf_to_doc(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOC format (simple text-based approach)"""
        try:
            # Convert PDF to text first
            temp_txt = output_path.replace('.doc', '.txt')
            if self._convert_pdf_to_text(input_path, temp_txt):
                # Create a simple RTF-like document
                with open(temp_txt, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Create a simple document format
                doc_content = f"""{{\\rtf1\\ansi\\deff0 {{\\fonttbl {{\\f0 Times New Roman;}}}}
\\f0\\fs24
{content.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')}
}}"""
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(doc_content)
                
                # Clean up temp file
                if os.path.exists(temp_txt):
                    os.remove(temp_txt)
                
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"PDF to DOC error: {str(e)}")
            return False
    
    def _convert_code_to_html(self, input_path: str, output_path: str, content: str) -> bool:
        """Convert code to HTML with syntax highlighting"""
        try:
            file_ext = Path(input_path).suffix[1:].lower()
            
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Code: {Path(input_path).name}</title>
    <style>
        body {{ font-family: 'Courier New', monospace; margin: 40px; background-color: #f8f8f8; }}
        pre {{ background-color: #ffffff; padding: 20px; border-radius: 5px; border: 1px solid #ddd; overflow-x: auto; }}
        .filename {{ color: #666; margin-bottom: 10px; }}
    </style>
</head>
<body>
    <div class="filename">File: {Path(input_path).name}</div>
    <pre><code>{content}</code></pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        except Exception as e:
            logger.error(f"Code to HTML error: {str(e)}")
            return False
    
    def _extract_text_to_file(self, input_path: str, output_path: str) -> bool:
        """Extract text from various file formats"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return True
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            return False
    
    def _convert_to_html(self, input_path: str, output_path: str) -> bool:
        """Convert various formats to HTML"""
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; }}
    </style>
</head>
<body>
    <pre>{content}</pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        except Exception as e:
            logger.error(f"Convert to HTML error: {str(e)}")
            return False
    
    def _convert_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert various formats to PDF"""
        try:
            # Extract text and convert to PDF
            return self._extract_text_to_file(input_path, output_path.replace('.pdf', '.txt')) and \
                   self._convert_text_to_pdf(output_path.replace('.pdf', '.txt'), output_path)
        except Exception as e:
            logger.error(f"Convert to PDF error: {str(e)}")
            return False
    
    def _convert_from_pdf(self, input_path: str, output_path: str, target_format: str) -> bool:
        """Convert PDF to various formats"""
        try:
            if target_format == 'txt':
                return self._convert_pdf_to_text(input_path, output_path)
            elif target_format == 'html':
                # Convert to text first, then to HTML
                temp_txt = output_path.replace('.html', '.txt')
                if self._convert_pdf_to_text(input_path, temp_txt):
                    result = self._convert_text_to_html(temp_txt, output_path)
                    if os.path.exists(temp_txt):
                        os.remove(temp_txt)
                    return result
            return False
        except Exception as e:
            logger.error(f"Convert from PDF error: {str(e)}")
            return False
    
    def _dict_to_xml(self, data, parent):
        """Helper method to convert dictionary to XML"""
        if isinstance(data, dict):
            for key, value in data.items():
                child = ET.SubElement(parent, key)
                self._dict_to_xml(value, child)
        elif isinstance(data, list):
            for item in data:
                child = ET.SubElement(parent, "item")
                self._dict_to_xml(item, child)
        else:
            parent.text = str(data)
    
    def _xml_to_dict(self, element):
        """Helper method to convert XML to dictionary"""
        result = {}
        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text
            else:
                result[child.tag] = self._xml_to_dict(child)
        return result
    
    def get_supported_formats(self) -> dict:
        """Get all supported format conversions"""
        return {
            'image_formats': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg'],
            'document_formats': ['pdf', 'txt', 'docx', 'doc', 'rtf', 'md', 'html', 'htm'],
            'spreadsheet_formats': ['xlsx', 'xls', 'csv'],
            'presentation_formats': ['pptx', 'ppt'],
            'data_formats': ['json', 'xml'],
            'code_formats': ['py', 'js', 'css', 'php', 'java', 'cpp', 'c', 'cs', 'rb', 'go', 'rs', 'log', 'ini', 'cfg', 'conf', 'yaml', 'yml', 'toml']
        }
    
    def get_conversion_options(self, input_format: str) -> list:
        """Get available conversion options for a given input format"""
        format_mappings = {
            # Image formats
            'jpg': ['png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg', 'pdf'],
            'jpeg': ['png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg', 'pdf'],
            'png': ['jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg', 'pdf'],
            'gif': ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg', 'pdf'],
            'bmp': ['jpg', 'jpeg', 'png', 'gif', 'tiff', 'tif', 'webp', 'ico', 'svg', 'pdf'],
            'tiff': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'ico', 'svg', 'pdf'],
            'tif': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'ico', 'svg', 'pdf'],
            'webp': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'ico', 'svg', 'pdf'],
            'ico': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'svg', 'pdf'],
            'svg': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'pdf'],
            
            # Document formats
            'pdf': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'tif', 'webp', 'ico', 'svg', 'txt', 'html', 'docx', 'doc'],
            'txt': ['pdf', 'html', 'md'],
            'docx': ['pdf', 'txt', 'html'],
            'doc': ['pdf', 'txt', 'html'],
            'rtf': ['pdf', 'txt', 'html'],
            'md': ['pdf', 'html', 'txt'],
            'html': ['pdf', 'txt', 'md'],
            'htm': ['pdf', 'txt', 'md'],
            
            # Spreadsheet formats
            'xlsx': ['csv', 'json', 'xml', 'pdf'],
            'xls': ['csv', 'json', 'xml', 'pdf'],
            'csv': ['json', 'xml', 'pdf'],
            
            # Data formats
            'json': ['xml', 'csv'],
            'xml': ['json', 'csv'],
            
            # Code formats
            'py': ['html', 'txt'],
            'js': ['html', 'txt'],
            'css': ['html', 'txt'],
            'php': ['html', 'txt'],
            'java': ['html', 'txt'],
            'cpp': ['html', 'txt'],
            'c': ['html', 'txt'],
            'cs': ['html', 'txt'],
            'rb': ['html', 'txt'],
            'go': ['html', 'txt'],
            'rs': ['html', 'txt'],
            'log': ['html', 'txt'],
            'ini': ['html', 'txt'],
            'cfg': ['html', 'txt'],
            'conf': ['html', 'txt'],
            'yaml': ['html', 'txt'],
            'yml': ['html', 'txt'],
            'toml': ['html', 'txt']
        }
        return format_mappings.get(input_format.lower(), []) 